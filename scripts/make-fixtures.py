# /// script
# requires-python = ">=3.10"
# dependencies = ["reportlab>=4.0", "python-docx>=1.1"]
# ///
"""make-fixtures.py — RAG検証用のPDF/DOCX fixtureを生成する。

fixtures/test-policy.txt と同じ設計思想:
一般常識と異なる固有の数値を含めることで、回答が「文書から取得されたもの」か
「モデルの推測（ハルシネーション）」かを判別できるようにする。

ファイルごとに扱う規程領域と固有値を変えてあり、複数ファイルを同一ワークスペースに
入れても「どのファイルから検索されたか」を質問で判別できる。

実行: uv run scripts/make-fixtures.py
出力: fixtures/test-expense.pdf, fixtures/test-attendance.docx
"""

from pathlib import Path

FIXTURES_DIR = Path(__file__).resolve().parent.parent / "fixtures"

PDF_TITLE = "LocalRAG 株式会社 経費規程（PDFパーステスト用サンプル文書）"
PDF_BODY = [
    "この文書は RAG パイプラインの PDF パース動作確認専用のサンプルです。実在の規程ではありません。",
    "",
    "第1条（出張日当）",
    "国内出張の日当は、1日あたり 3,400 円とする。",
    "",
    "第2条（交通費精算）",
    "交通費の精算申請は、利用日から 45 日以内に行わなければならない。",
    "",
    "第3条（会議費）",
    "1人あたりの会議費の上限は、1回につき 6,800 円とする。",
    "",
    "第4条（承認フロー）",
    "経費申請は金額が 87,000 円を超える場合、部門長に加えて経理部長の承認を要する。",
]

DOCX_TITLE = "LocalRAG 株式会社 勤怠規程（DOCXパーステスト用サンプル文書）"
DOCX_BODY = [
    "この文書は RAG パイプラインの DOCX パース動作確認専用のサンプルです。実在の規程ではありません。",
    "",
    "第1条（コアタイム）",
    "フレックスタイム制のコアタイムは、午前 10 時 20 分から午後 3 時 40 分までとする。",
    "",
    "第2条（残業の事前申請）",
    "残業は 1 日 2.5 時間を超える場合、前営業日までに申請しなければならない。",
    "",
    "第3条（在宅勤務）",
    "在宅勤務は週に 3 日まで認められる。申請は内線番号 8912 の総務課へ行う。",
    "",
    "第4条（休憩時間）",
    "6 時間を超える勤務では、73 分の休憩を取得しなければならない。",
]


def make_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
    c = canvas.Canvas(str(path), pagesize=A4)
    _, height = A4
    y = height - 60
    c.setFont("HeiseiKakuGo-W5", 14)
    c.drawString(50, y, PDF_TITLE)
    y -= 30
    c.setFont("HeiseiKakuGo-W5", 11)
    for line in PDF_BODY:
        c.drawString(50, y, line)
        y -= 18
    c.save()


def make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(DOCX_TITLE, level=1)
    for line in DOCX_BODY:
        doc.add_paragraph(line)
    doc.save(str(path))


if __name__ == "__main__":
    FIXTURES_DIR.mkdir(exist_ok=True)
    pdf_path = FIXTURES_DIR / "test-expense.pdf"
    docx_path = FIXTURES_DIR / "test-attendance.docx"
    make_pdf(pdf_path)
    make_docx(docx_path)
    print(f"created: {pdf_path}")
    print(f"created: {docx_path}")
